import os
import logging
from typing import Optional

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pypdf
    """
    if PdfReader is None:
        raise Exception("pypdf not available. Please install the pypdf package.")
    
    try:
        text = ""
        reader = PdfReader(file_path)
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        text = text.strip()
        
        if not text:
            raise Exception("No text could be extracted from the PDF")
        
        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text
        
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx
    """
    if Document is None:
        raise Exception("python-docx not available. Please install python-docx package.")
    
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        text = text.strip()
        
        if not text:
            raise Exception("No text could be extracted from the DOCX file")
        
        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from file based on its extension
    """
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        logger.error(f"File processing error for {file_path}: {str(e)}")
        raise

def get_file_info(file_path: str) -> dict:
    """
    Get basic information about a file
    """
    try:
        stat = os.stat(file_path)
        return {
            'size': stat.st_size,
            'extension': os.path.splitext(file_path)[1].lower(),
            'name': os.path.basename(file_path)
        }
    except Exception as e:
        logger.error(f"Error getting file info for {file_path}: {str(e)}")
        return {}
